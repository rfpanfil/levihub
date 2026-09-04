# arquivo: api/routers/transpor.py
# Lógica musical de transposição de acordes + endpoints do transpositor.

import re
import io
import docx
from fastapi import APIRouter, File, Form, UploadFile

from api.models import (
    TransposeCifraRequest, TransposeCifraResponse,
    TransposeSequenceRequest, TransposeSequenceResponse,
)

router = APIRouter(tags=["Transpositor"])


# =============================================================================
# LÓGICA MUSICAL (funções puras — sem I/O)
# =============================================================================

MAPA_NOTAS = {
    "C": 0, "C#": 1, "Db": 1, "D": 2, "D#": 3, "Eb": 3, "E": 4, "F": 5,
    "F#": 6, "Gb": 6, "G": 7, "G#": 8, "Ab": 8, "A": 9, "A#": 10, "Bb": 10, "B": 11,
    "E#": 5, "B#": 0, "Fb": 4, "Cb": 11,
}
MAPA_VALORES_NOTAS = ["C", "C#", "D", "D#", "E", "F", "F#", "G", "G#", "A", "A#", "B"]

EXPLICACAO_TEORICA = {
    "E#": "Mi sustenido (E#) é enarmônica de Fá (F).",
    "B#": "Si sustenido (B#) é enarmônica de Dó (C).",
    "Fb": "Fá bemol (Fb) é enarmônica de Mi (E).",
    "Cb": "Dó bemol (Cb) é enarmônica de Si (B).",
}


def transpor_nota_individual(nota_str: str, semitons: int) -> str:
    nota_key = next(
        (key for key in MAPA_NOTAS if key.lower() == nota_str.lower()), None
    )
    if not nota_key:
        return nota_str
    valor_original = MAPA_NOTAS[nota_key]
    novo_valor = (valor_original + semitons + 12) % 12
    return MAPA_VALORES_NOTAS[novo_valor]


def normalizar_nota(nota_str: str, explicacoes_set=None) -> str:
    if nota_str.endswith("##"):
        base = nota_str.replace("##", "")
        base_key = next((k for k in MAPA_NOTAS if k.lower() == base.lower()), None)
        if base_key is not None:
            novo_valor = (MAPA_NOTAS[base_key] + 2) % 12
            nova_nota = MAPA_VALORES_NOTAS[novo_valor]
            if explicacoes_set is not None:
                explicacoes_set.add(f"A nota {nota_str} é enarmônica de {nova_nota} (Duplo Sustenido).")
            return nova_nota

    if nota_str.endswith("bb"):
        base = nota_str.replace("bb", "")
        base_key = next((k for k in MAPA_NOTAS if k.lower() == base.lower()), None)
        if base_key is not None:
            novo_valor = (MAPA_NOTAS[base_key] - 2 + 12) % 12
            nova_nota = MAPA_VALORES_NOTAS[novo_valor]
            if explicacoes_set is not None:
                explicacoes_set.add(f"A nota {nota_str} é enarmônica de {nova_nota} (Duplo Bemol).")
            return nova_nota

    return nota_str


def transpor_acordes_sequencia(acordes_originais, acao, intervalo):
    semitons = int(intervalo * 2) * (1 if acao == "Aumentar" else -1)
    acordes_transpostos = []
    explicacoes_entrada = set()

    for acorde_original in acordes_originais:
        match = re.match(r"^([A-G](?:##|bb|#|b)?)(.*)", acorde_original, re.IGNORECASE)
        if not match:
            acordes_transpostos.append(f"{acorde_original}?")
            continue

        nota_bruta, resto = match.groups()
        nota_fundamental = normalizar_nota(nota_bruta, explicacoes_entrada)

        if nota_fundamental == nota_bruta:
            nota_key = next(
                (k for k in EXPLICACAO_TEORICA if k.lower() == nota_fundamental.lower()), None
            )
            if nota_key:
                explicacoes_entrada.add(EXPLICACAO_TEORICA[nota_key])

        nova_fundamental = transpor_nota_individual(nota_fundamental, semitons)

        if "/" in resto:
            partes = resto.split("/")
            qualidade = partes[0]
            baixo_norm = normalizar_nota(partes[1], explicacoes_entrada)
            novo_baixo = transpor_nota_individual(baixo_norm, semitons)
            acorde_final = f"{nova_fundamental}{qualidade}/{novo_baixo}"
        else:
            acorde_final = f"{nova_fundamental}{resto}"

        acordes_transpostos.append(acorde_final)

    return acordes_transpostos, list(explicacoes_entrada)


def is_chord_line(line: str) -> bool:
    line = line.strip()
    if not line:
        return False
    chord_pattern = re.compile(
        r"^\(?[A-G](?:##|bb|#|b)?(?:m|M|dim|aug|sus|add|maj|º|°)?(?:2|4|5|6|7|9|11|13)?(?:\([^)]+\))?(?:/[A-G](?:##|bb|#|b)?)?\)?$"
    )
    # Removemos pontuações grudadas, mas não apagamos parênteses para não estragar (11)
    clean_line = re.sub(r'[,.]', '', line.strip())
    words = clean_line.split()
    
    if not words:
        return False
        
    markers = {"1ª:", "2ª:", "3ª:", "1x", "2x", "3x", "4x", "intro:", "tom:", "final:", "refrão", "refrão:", "ponte", "ponte:", ":/", "/:", "-", "|"}
    valid_words = [w for w in words if w.lower() not in markers]
    
    if not valid_words:
        return False
        
    chord_count = sum(1 for word in valid_words if chord_pattern.match(word))
    return (chord_count / len(valid_words)) > 0.5


def processar_cifra(texto_cifra: str, acao: str, intervalo: float) -> str:
    semitons = int(intervalo * 2) * (1 if acao == "Aumentar" else -1)
    padrao_acorde = r"(^|[\s(])([A-G](?:##|bb|#|b)?(?:m|M|dim|aug|sus|add|maj|º|°)?(?:2|4|5|6|7|9|11|13)?(?:\([^)]+\))?(?:/[A-G](?:##|bb|#|b)?)?)(?=\s|$|\)|,)"

    def replacer(match):
        prefixo, acorde = match.groups()
        prefixo = prefixo or ""
        
        def transpose_part(part):
            m = re.match(r"^([A-G](?:##|bb|#|b)?)(.*)", part)
            if m:
                root, quality = m.groups()
                root_trans = transpor_nota_individual(normalizar_nota(root), semitons)
                return f"{root_trans}{quality}"
            return part

        if "/" in acorde:
            base, baixo = acorde.split("/", 1)
            novo_acorde = f"{transpose_part(base)}/{transpose_part(baixo)}"
        else:
            novo_acorde = transpose_part(acorde)
            
        return f"{prefixo}{novo_acorde}"

    def processar_linha(linha: str) -> str:
        # Fatiamento inteligente: separa a linha em blocos por Tab ou 3+ espaços seguidos
        partes = re.split(r'(\s{3,}|\t+)', linha)
        if len(partes) == 1:
            # Avalia a linha toda normalmente
            return re.sub(padrao_acorde, replacer, linha) if is_chord_line(linha) else linha
        
        # Avalia cada bloco fatiado individualmente
        nova_linha = []
        for p in partes:
            if p.isspace():
                nova_linha.append(p)
            elif is_chord_line(p):
                nova_linha.append(re.sub(padrao_acorde, replacer, p))
            else:
                nova_linha.append(p)
        return "".join(nova_linha)

    linhas = texto_cifra.split("\n")
    linhas_finais = [processar_linha(linha) for linha in linhas]
    return "\n".join(linhas_finais)


async def ler_conteudo_arquivo(file: UploadFile) -> str:
    content = await file.read()
    if file.filename.endswith(".docx"):
        try:
            doc = docx.Document(io.BytesIO(content))
            return "\n".join([p.text for p in doc.paragraphs])
        except Exception as e:
            return f"Erro ao ler arquivo .docx: {str(e)}"
    return content.decode("utf-8")


# =============================================================================
# ENDPOINTS
# =============================================================================

@router.post("/transpose-sequence", response_model=TransposeSequenceResponse, dependencies=[])
async def transpose_sequence_endpoint(request: TransposeSequenceRequest):
    transposed, expl = transpor_acordes_sequencia(
        request.chords, request.action, request.interval
    )
    return {
        "original_chords": request.chords,
        "transposed_chords": transposed,
        "explanations": expl,
    }


@router.post("/transpose-text", response_model=TransposeCifraResponse, dependencies=[])
async def transpose_text_endpoint(request: TransposeCifraRequest):
    res = processar_cifra(request.cifra_text, request.action, request.interval)
    return {"transposed_cifra": res}


from fastapi.responses import StreamingResponse

def get_run_style(run):
    return (
        run.bold,
        run.italic,
        run.underline,
        run.font.name,
        run.font.size,
        run.font.color.rgb if run.font.color else None
    )

def set_run_text_safe(run, new_text):
    # Acessa os nós de texto diretamente sem apagar os nós de desenho (shapes, brackets, imagens)
    t_elements = run._r.xpath('.//w:t')
    if t_elements:
        t_elements[0].text = new_text
        for t in t_elements[1:]:
            t.text = ""
    else:
        if new_text:
            from docx.oxml import OxmlElement
            t = OxmlElement('w:t')
            t.text = new_text
            run._r.append(t)

def has_special_content(run):
    xml = run._element.xml
    tags = ["<w:drawing", "<w:pict", "<mc:AlternateContent", "v:shape", "<w:sym", "<w:object"]
    return any(tag in xml for tag in tags)

@router.post("/transpose-file", dependencies=[])
async def transpose_file_endpoint(
    file: UploadFile = File(...),
    action: str = Form(...),
    interval: float = Form(...),
):
    content = await file.read()
    
    if file.filename.endswith(".docx"):
        try:
            doc = docx.Document(io.BytesIO(content))
            
            for paragraph in doc.paragraphs:
                if not paragraph.text.strip():
                    continue
                    
                # Fusão (Merge) de runs idênticos para evitar que palavras sejam fatiadas no meio 
                # por corretores ortográficos invisíveis do Word (ex: "Em" -> "E" + "m").
                runs = paragraph.runs
                if not runs:
                    continue
                    
                merged_runs = []
                current_run = runs[0]
                current_style = get_run_style(current_run)
                current_text = current_run.text
                
                for run in runs[1:]:
                    style = get_run_style(run)
                    # Não mesclamos runs se algum deles possuir conteúdo especial (símbolo, seta, desenho)
                    if style == current_style and not has_special_content(run) and not has_special_content(current_run):
                        current_text += run.text
                        set_run_text_safe(run, "")
                    else:
                        set_run_text_safe(current_run, current_text)
                        merged_runs.append(current_run)
                        current_run = run
                        current_style = style
                        current_text = current_run.text
                set_run_text_safe(current_run, current_text)
                merged_runs.append(current_run)
                    
                # Aplicamos a transposição run por run (já unificados)
                for run in merged_runs:
                    original_text = run.text
                    if not original_text.strip():
                        continue
                    
                    new_run_text = processar_cifra(original_text, action, interval)
                    if new_run_text != original_text:
                        set_run_text_safe(run, new_run_text)
                        
            # Salva o arquivo em buffer de memória
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            headers = {
                "Content-Disposition": f'attachment; filename="Transposto_{file.filename}"'
            }
            return StreamingResponse(
                buffer, 
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers=headers
            )
        except Exception as e:
            return {"transposed_cifra": f"Erro interno DOCX: {str(e)}"}
            
    # Tratamento normal de .txt
    texto = content.decode("utf-8")
    res = processar_cifra(texto, action, interval)
    return {"transposed_cifra": res}
